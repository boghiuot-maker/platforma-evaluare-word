from flask import Flask, request, jsonify, render_template, send_from_directory, redirect, url_for
from werkzeug.utils import secure_filename
from docx import Document
import os, csv, io, datetime
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

UPLOAD_FOLDER = 'uploads'
DATA_FOLDER = 'data'
RESULTS_CSV = os.path.join(DATA_FOLDER, 'results.csv')
TEMPLATE_FILE = 'test_practic_suport.docx'

ALLOWED_EXTENSIONS = {'doc', 'docx'}

app = Flask(__name__, static_folder='static', template_folder='templates')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = os.environ.get('FLASK_SECRET', 'schimba_secretul')

ADMIN_PASS = os.environ.get('ADMIN_PASS', 'profesor')

CRITERIA = {
    'titlu_corect': 2,
    'are_imagine': 1,
    'are_tabel': 1,
    'are_antet': 1,
    'are_pie_pagina': 1,
    'are_lista': 1,
    'numar_total_cuvinte': 3
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def verifica_titlu(doc):
    for p in doc.paragraphs:
        style = getattr(p.style, 'name', '').lower()
        if style.startswith('heading'):
            return True
    return False

def verifica_imagine(doc):
    return bool(getattr(doc, 'inline_shapes', []))

def verifica_tabel(doc):
    return len(doc.tables) > 0

def verifica_antet(doc):
    try:
        for sec in doc.sections:
            if sec.header and any(p.text.strip() for p in sec.header.paragraphs):
                return True
    except:
        return False
    return False

def verifica_pie_pagina(doc):
    try:
        for sec in doc.sections:
            if sec.footer and any(p.text.strip() for p in sec.footer.paragraphs):
                return True
    except:
        return False
    return False

def verifica_lista(doc):
    for p in doc.paragraphs:
        style = getattr(p.style, 'name', '').lower()
        if 'list' in style or 'bullet' in style or 'number' in style:
            return True
    return False

def numar_cuvinte(doc):
    total = 0
    for p in doc.paragraphs:
        total += len(p.text.split())
    return total

def genereaza_feedback(rezultate, nr_cuv):
    fb = []
    if not rezultate.get('titlu_corect'): fb.append('Titlul nu este setat ca Heading.')
    if not rezultate.get('are_imagine'): fb.append('Nu ai inserat o imagine.')
    if not rezultate.get('are_tabel'): fb.append('Nu ai inserat un tabel.')
    if not rezultate.get('are_antet'): fb.append('Antetul lipsește.')
    if not rezultate.get('are_pie_pagina'): fb.append('Nu există text în subsol.')
    if not rezultate.get('are_lista'): fb.append('Nu ai folosit liste (bullet/numere).')
    if nr_cuv < 100: fb.append('Numărul de cuvinte este prea mic (<100).')
    return fb if fb else ['Excelent! Toate cerințele au fost îndeplinite.']

def calculeaza_scor(rezultate, nr_cuv):
    scor = 0
    for crit, pts in CRITERIA.items():
        if crit == 'numar_total_cuvinte':
            if nr_cuv >= 100: scor += pts
        else:
            if rezultate.get(crit): scor += pts
    return scor, sum(CRITERIA.values())

def append_result(row):
    header = ['timestamp','nume_elev','clasa','data_test','fisier','scor','punctaj_maxim','nr_cuvinte'] + list(CRITERIA.keys())
    exists = os.path.exists(RESULTS_CSV)
    with open(RESULTS_CSV, 'a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        if not exists:
            writer.writerow(header)
        writer.writerow([row.get(h,'') for h in header])

def already_submitted(name, clasa, data_test):
    if not os.path.exists(RESULTS_CSV):
        return False
    df = pd.read_csv(RESULTS_CSV)
    if 'nume_elev' in df.columns and 'clasa' in df.columns and 'data_test' in df.columns:
        same = df[(df['nume_elev']==name) & (df['clasa']==clasa) & (df['data_test']==data_test)]
        return not same.empty
    return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/download-template')
def download_template():
    return send_from_directory(DATA_FOLDER, TEMPLATE_FILE, as_attachment=True)

@app.route('/evaluate', methods=['POST'])
def evaluate():
    if 'file' not in request.files:
        return jsonify({'error':'Nu ai încărcat fișier'}),400
    file = request.files['file']
    name = request.form.get('nume','').strip()
    clasa = request.form.get('clasa','').strip()
    data_test = request.form.get('data_test','').strip()
    if not name:
        return jsonify({'error':'Numele este obligatoriu'}),400
    if not clasa:
        return jsonify({'error':'Clasa este obligatorie'}),400
    if not data_test:
        return jsonify({'error':'Data testului este obligatorie'}),400
    if already_submitted(name, clasa, data_test):
        return jsonify({'error':'Ai deja o evaluare pentru această clasă și dată'}),400
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error':'Fișier invalid'}),400
    filename = secure_filename(f"{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}")
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    file.save(os.path.join(UPLOAD_FOLDER, filename))
    doc = Document(os.path.join(UPLOAD_FOLDER, filename))
    rezultate = {
        'titlu_corect': verifica_titlu(doc),
        'are_imagine': verifica_imagine(doc),
        'are_tabel': verifica_tabel(doc),
        'are_antet': verifica_antet(doc),
        'are_pie_pagina': verifica_pie_pagina(doc),
        'are_lista': verifica_lista(doc)
    }
    nr = numar_cuvinte(doc)
    scor, max_total = calculeaza_scor(rezultate, nr)
    feedback = genereaza_feedback(rezultate, nr)
    row = {
        'timestamp': datetime.datetime.now().isoformat(),
        'nume_elev': name,
        'clasa': clasa,
        'data_test': data_test,
        'fisier': filename,
        'scor': scor,
        'punctaj_maxim': max_total,
        'nr_cuvinte': nr
    }
    for k in CRITERIA.keys():
        row[k] = rezultate.get(k, False) if k != 'numar_total_cuvinte' else (nr>=100)
    append_result(row)
    return jsonify({'rezultate':rezultate,'scor':scor,'punctaj_maxim':max_total,'feedback':feedback,'nr_cuvinte':nr})

@app.route('/admin', methods=['GET','POST'])
def admin():
    pw = request.args.get('pw') or request.form.get('pw')
    if request.method == 'POST' and pw == ADMIN_PASS:
        return redirect(url_for('admin') + f"?pw={ADMIN_PASS}")
    if pw != ADMIN_PASS:
        return render_template('admin_login.html', error=None)
    if os.path.exists(RESULTS_CSV):
        df = pd.read_csv(RESULTS_CSV)
        records = df.to_dict(orient='records')
    else:
        records = []
    return render_template('admin.html', records=records)

@app.route('/admin/export/csv')
def export_csv():
    if request.args.get('pw') != ADMIN_PASS:
        return 'Unauthorized', 403
    if not os.path.exists(RESULTS_CSV):
        return 'No results', 404
    return send_from_directory(DATA_FOLDER, 'results.csv', as_attachment=True)


@app.route('/admin/export/xlsx')
def export_xlsx():
    if request.args.get('pw') != ADMIN_PASS:
        return 'Unauthorized', 403
    if not os.path.exists(RESULTS_CSV):
        return 'No results', 404
    df = pd.read_csv(RESULTS_CSV)
    # Add per-criteria points columns
    for crit, pts in CRITERIA.items():
        col_name = f"{crit}_pts"
        if crit == 'numar_total_cuvinte':
            # if numeric in nr_cuvinte, award pts if >=100
            if 'nr_cuvinte' in df.columns:
                df[col_name] = df['nr_cuvinte'].apply(lambda v, p=pts: p if str(v).isdigit() and int(str(v))>=100 else 0)
            else:
                df[col_name] = 0
        else:
            if crit in df.columns:
                df[col_name] = df[crit].apply(lambda v, p=pts: p if str(v).lower() in ['true','1','t','y'] else 0)
            else:
                df[col_name] = 0
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='results')
    out.seek(0)
    return (out.read(), 200, {
        'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'Content-Disposition': 'attachment; filename="results_with_points.xlsx"'
    })

@app.route('/admin/result/<filename>/pdf')
def result_pdf(filename):
    if request.args.get('pw') != ADMIN_PASS:
        return 'Unauthorized', 403
    if not os.path.exists(RESULTS_CSV):
        return 'No results', 404
    df = pd.read_csv(RESULTS_CSV)
    row = df[df['fisier'] == filename]
    if row.empty:
        return 'Not found', 404
    row = row.iloc[0].to_dict()
    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont('Helvetica-Bold', 14)
    c.drawString(50, y, f"Rezultat evaluare - {row.get('nume_elev')}")
    y -= 30
    c.setFont('Helvetica', 11)
    main_fields = ['timestamp','nume_elev','clasa','data_test','fisier','scor','punctaj_maxim','nr_cuvinte']
    for k in main_fields:
        c.drawString(50, y, f"{k}: {row.get(k,'')}")
        y -= 16
    y -= 8
    c.setFont('Helvetica-Bold', 12)
    c.drawString(50, y, 'Detalii:')
    y -= 18
    c.setFont('Helvetica', 11)
    for crit in list({'titlu_corect','are_imagine','are_tabel','are_antet','are_pie_pagina','are_lista','numar_total_cuvinte'}):
        c.drawString(60, y, f"- {crit}: {row.get(crit,'')}")
        y -= 14
        if y < 60:
            c.showPage()
            y = height - 50
    c.showPage()
    c.save()
    out.seek(0)
    return (out.read(), 200, {
        'Content-Type': 'application/pdf',
        'Content-Disposition': f'attachment; filename="{filename}_result.pdf"'
    })

if __name__ == '__main__':
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(DATA_FOLDER, exist_ok=True)
    app.run(host='0.0.0.0', port=5000, debug=True)



@app.route('/admin/delete_all', methods=['POST'])
def admin_delete_all():
    if request.args.get('pw') != ADMIN_PASS and request.form.get('pw') != ADMIN_PASS:
        return 'Unauthorized', 403
    # delete results csv
    if os.path.exists(RESULTS_CSV):
        try:
            os.remove(RESULTS_CSV)
        except Exception:
            pass
    # clear uploads folder
    if os.path.exists(UPLOAD_FOLDER):
        for fname in os.listdir(UPLOAD_FOLDER):
            fpath = os.path.join(UPLOAD_FOLDER, fname)
            try:
                if os.path.isfile(fpath):
                    os.remove(fpath)
            except Exception:
                pass
    return redirect(f"/admin?pw={ADMIN_PASS}")
